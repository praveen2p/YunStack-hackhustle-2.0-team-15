from __future__ import annotations

import json
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from app.core.config import settings

try:
    from google import genai
except ImportError:  # pragma: no cover
    genai = None

try:
    import cv2
except ImportError:  # pragma: no cover
    cv2 = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:  # pragma: no cover
    RapidOCR = None


@dataclass
class ProcessedDocument:
    title: str
    content: str
    summary: str
    extracted_data: dict[str, Any]
    record_type: str
    risk_label: str
    risk_confidence: str


# Intentionally disabled: the previous XGBoost classifier was trained on only
# five synthetic rows, which made the output misleading for clinical UX.
RISK_MODEL = None
RISK_ENGINE_NAME = "Clinical threshold heuristic"
OCR_ENGINE = RapidOCR() if RapidOCR is not None else None


def _clean_json(text: str) -> dict[str, Any]:
    cleaned = re.sub(r"```json|```", "", text).strip()
    return json.loads(cleaned)


def _fill_missing_features(extracted: dict[str, Any]) -> list[float]:
    return [
        extracted.get("age") or 50,
        extracted.get("glucose") or 120,
        extracted.get("bp_sys") or 120,
        extracted.get("bp_dia") or 80,
        extracted.get("hba1c") or 5.5,
        extracted.get("bmi") or 25,
        extracted.get("cholesterol") or 180,
    ]


def _predict_risk(values: list[float]) -> tuple[str, str]:
    _, glucose, bp_sys, bp_dia, hba1c, bmi, cholesterol = values
    high_signals = [
        glucose >= 200,
        hba1c >= 8.0,
        bp_sys >= 160,
        bp_dia >= 100,
        cholesterol >= 240,
        bmi >= 35,
    ]
    moderate_signals = [
        glucose >= 126,
        hba1c >= 6.5,
        bp_sys >= 140,
        bp_dia >= 90,
        cholesterol >= 200,
        bmi >= 30,
    ]

    if any(high_signals) or sum(moderate_signals) >= 3:
        return "High", "Clinical thresholds"
    if any(moderate_signals):
        return "Medium", "Clinical thresholds"
    return "Low", "Clinical thresholds"


def _format_value(value: Any) -> str:
    if value is None:
        return "Not reported"
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    if isinstance(value, list):
        return ", ".join(str(item) for item in value) if value else "None noted"
    return str(value)


def _build_readable_content(
    summary: str,
    record_type: str,
    risk_label: str,
    risk_confidence: str,
    extracted: dict[str, Any],
    filename: str,
) -> str:
    lines = [
        f"Summary: {summary}",
        f"Document type: {record_type}",
        f"Risk level: {risk_label}",
        f"Confidence: {risk_confidence}",
        f"Source file: {filename}",
        "",
        "Clinical details",
        f"Age: {_format_value(extracted.get('age'))}",
        f"Glucose: {_format_value(extracted.get('glucose'))}",
        f"Blood pressure (systolic): {_format_value(extracted.get('bp_sys'))}",
        f"Blood pressure (diastolic): {_format_value(extracted.get('bp_dia'))}",
        f"HbA1c: {_format_value(extracted.get('hba1c'))}",
        f"Cholesterol: {_format_value(extracted.get('cholesterol'))}",
        f"BMI: {_format_value(extracted.get('bmi'))}",
        f"Diagnoses: {_format_value(extracted.get('diagnosis'))}",
        f"Medications: {_format_value(extracted.get('medications'))}",
    ]
    return "\n".join(lines)


def _normalize_extracted_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned_lines = [line for line in lines if line]
    return "\n".join(cleaned_lines)


def _extract_text_with_ocr(image_source: str | bytes) -> str:
    if OCR_ENGINE is None:
        return ""

    try:
        result, _ = OCR_ENGINE(image_source)
    except Exception:
        return ""

    if not result:
        return ""

    segments: list[str] = []
    for item in result:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            text = item[1]
            if isinstance(text, str) and text.strip():
                segments.append(text.strip())

    return _normalize_extracted_text("\n".join(segments))


def _extract_with_gemini_from_image(image_path: str) -> dict[str, Any] | None:
    if genai is None or Image is None or not settings.effective_gemini_api_key:
        return None

    client = genai.Client(api_key=settings.effective_gemini_api_key)
    image = Image.open(image_path)
    prompt = """
You are a medical AI system.

Extract structured clinical data AND provide a summary.

Return ONLY JSON:

{
  "age": number,
  "glucose": number,
  "bp_sys": number,
  "bp_dia": number,
  "hba1c": number,
  "cholesterol": number,
  "bmi": number,
  "diagnosis": [],
  "medications": [],
  "document_type": "lab|prescription|imaging|consultation|vaccination|discharge",
  "summary": "short clinical summary"
}

If any value is missing, return null.
Choose the best document_type from the allowed values based on visible content.
Do not add explanation.
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[prompt, image],
    )
    return _clean_json(response.text)


def _extract_with_gemini_from_text(text: str) -> dict[str, Any] | None:
    if genai is None or not settings.effective_gemini_api_key:
        return None

    client = genai.Client(api_key=settings.effective_gemini_api_key)
    prompt = f"""
You are a medical AI system.

Extract structured clinical data AND provide a summary from the text below.

Return ONLY JSON:

{{
  "age": number,
  "glucose": number,
  "bp_sys": number,
  "bp_dia": number,
  "hba1c": number,
  "cholesterol": number,
  "bmi": number,
  "diagnosis": [],
  "medications": [],
  "document_type": "lab|prescription|imaging|consultation|vaccination|discharge",
  "summary": "short clinical summary"
}}

If any value is missing, return null.
Choose the best document_type from the allowed values based on the text.
Do not add explanation.

Medical text:
{text[:12000]}
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )
    return _clean_json(response.text)


def _preprocess_image_bytes(file_bytes: bytes, suffix: str) -> str | None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as raw_file:
        raw_file.write(file_bytes)
        raw_path = raw_file.name

    if cv2 is None:
        return raw_path

    image = cv2.imread(raw_path)
    if image is None:
        return raw_path

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as processed_file:
        cv2.imwrite(processed_file.name, blur)
        return processed_file.name


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    if fitz is None:
        return ""
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        return "\n".join(page.get_text() for page in document)


def _extract_text_from_scanned_pdf(file_bytes: bytes) -> str:
    if fitz is None:
        return ""

    page_text: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            page_text.append(_extract_text_with_ocr(pixmap.tobytes("png")))
    return _normalize_extracted_text("\n\n".join(page_text))


def _extract_text_from_docx(temp_path: Path) -> str:
    if Document is None:
        return ""
    document = Document(temp_path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def _fallback_extract_from_text(text: str, filename: str) -> dict[str, Any]:
    lowered = text.lower()

    def _extract_number(pattern: str) -> float | None:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        return float(match.group(1)) if match else None

    diagnoses = []
    if "diabetes" in lowered:
        diagnoses.append("diabetes")
    if "hypertension" in lowered or "blood pressure" in lowered:
        diagnoses.append("hypertension")

    medications = re.findall(r"(metformin|lisinopril|insulin|amlodipine)", lowered)
    summary = text.strip().splitlines()[0] if text.strip() else f"Clinical document uploaded from {filename}."

    return {
        "age": _extract_number(r"age[:\s]+(\d+)"),
        "glucose": _extract_number(r"glucose[:\s]+(\d+(?:\.\d+)?)"),
        "bp_sys": _extract_number(r"bp[_\s-]*sys[:\s]+(\d+(?:\.\d+)?)"),
        "bp_dia": _extract_number(r"bp[_\s-]*dia[:\s]+(\d+(?:\.\d+)?)"),
        "hba1c": _extract_number(r"hba1c[:\s]+(\d+(?:\.\d+)?)"),
        "cholesterol": _extract_number(r"cholesterol[:\s]+(\d+(?:\.\d+)?)"),
        "bmi": _extract_number(r"bmi[:\s]+(\d+(?:\.\d+)?)"),
        "diagnosis": diagnoses,
        "medications": medications,
        "document_type": _classify_document_type(text, filename),
        "summary": summary[:500],
    }


DOCUMENT_TYPES = {"lab", "prescription", "imaging", "consultation", "vaccination", "discharge"}


def _score_keywords(haystack: str, keywords: list[str]) -> int:
    return sum(1 for keyword in keywords if re.search(rf"\b{re.escape(keyword)}\b", haystack))


def _classify_document_type(text: str, filename: str) -> str:
    haystack = f"{Path(filename).stem} {text}".lower()

    scores = {
        "prescription": _score_keywords(haystack, [
            "prescription", "rx", "prescribed", "tablet", "capsule", "syrup", "ointment",
            "dose", "dosage", "frequency", "medication", "medications", "pharmacy",
            "refill", "take", "oral",
        ]),
        "lab": _score_keywords(haystack, [
            "lab", "laboratory", "pathology", "test", "tests", "result", "results",
            "report", "specimen", "sample", "glucose", "hba1c", "cholesterol", "cbc",
            "hemoglobin", "platelet", "creatinine", "bilirubin", "urine", "blood",
            "reference range",
        ]),
        "imaging": _score_keywords(haystack, [
            "xray", "x-ray", "mri", "ct", "scan", "ultrasound", "sonography",
            "radiology", "imaging", "contrast", "impression", "findings",
        ]),
        "discharge": _score_keywords(haystack, [
            "discharge", "admission", "admitted", "hospital course", "condition on discharge",
            "discharge summary", "ward", "inpatient",
        ]),
        "vaccination": _score_keywords(haystack, [
            "vaccination", "vaccine", "immunization", "immunisation", "booster",
            "covid", "influenza", "hepatitis", "mmr", "tdap",
        ]),
        "consultation": _score_keywords(haystack, [
            "consultation", "consult", "clinical note", "progress note", "chief complaint",
            "assessment", "plan", "follow up", "follow-up", "diagnosis", "symptoms",
            "examination",
        ]),
    }

    # Medication-heavy documents are prescriptions even when they mention a diagnosis.
    if scores["prescription"] >= 2:
        return "prescription"
    if scores["discharge"] >= 2:
        return "discharge"
    if scores["vaccination"] >= 1:
        return "vaccination"
    if scores["imaging"] >= 2:
        return "imaging"
    if scores["lab"] >= 2:
        return "lab"
    if scores["consultation"] >= 2:
        return "consultation"

    best_type, best_score = max(scores.items(), key=lambda item: item[1])
    return best_type if best_score > 0 else "consultation"


def _normalize_document_type(value: Any, text: str, filename: str) -> str:
    if isinstance(value, str):
        normalized = value.strip().lower().replace(" ", "_")
        aliases = {
            "laboratory": "lab",
            "lab_report": "lab",
            "diagnostic": "lab",
            "diagnostic_report": "lab",
            "radiology": "imaging",
            "scan": "imaging",
            "immunization": "vaccination",
            "immunisation": "vaccination",
            "clinical_note": "consultation",
            "consult": "consultation",
            "discharge_summary": "discharge",
        }
        normalized = aliases.get(normalized, normalized)
        if normalized in DOCUMENT_TYPES:
            return normalized
    return _classify_document_type(text, filename)


def _extract_structured_data_from_text(text: str, filename: str) -> dict[str, Any]:
    normalized_text = _normalize_extracted_text(text)
    if not normalized_text:
        return _fallback_extract_from_text("", filename)

    extracted = _extract_with_gemini_from_text(normalized_text)
    if extracted is not None:
        extracted["document_type"] = _normalize_document_type(extracted.get("document_type"), normalized_text, filename)
        return extracted
    return _fallback_extract_from_text(normalized_text, filename)


def process_uploaded_document(file_bytes: bytes, filename: str, content_type: str | None) -> ProcessedDocument:
    suffix = Path(filename).suffix.lower() or ".bin"
    extracted_text = ""

    if content_type and content_type.startswith("image/"):
        processed_path = _preprocess_image_bytes(file_bytes, suffix)
        if processed_path:
            extracted_text = _extract_text_with_ocr(processed_path)
            if not extracted_text:
                extracted = _extract_with_gemini_from_image(processed_path)
            else:
                extracted = _extract_structured_data_from_text(extracted_text, filename)
        else:
            extracted = None
    elif suffix == ".pdf":
        extracted_text = _extract_text_from_pdf(file_bytes)
        if len(extracted_text.strip()) < 40:
            extracted_text = _extract_text_from_scanned_pdf(file_bytes) or extracted_text
        extracted = _extract_structured_data_from_text(extracted_text, filename) if extracted_text else None
    elif suffix == ".docx":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx.write(file_bytes)
            extracted_text = _extract_text_from_docx(Path(temp_docx.name))
        extracted = _extract_structured_data_from_text(extracted_text, filename) if extracted_text else None
    else:
        extracted_text = file_bytes.decode("utf-8", errors="ignore")
        extracted = _extract_structured_data_from_text(extracted_text, filename) if extracted_text.strip() else None

    if extracted is None:
        extracted = _fallback_extract_from_text(extracted_text, filename)

    record_type = _normalize_document_type(extracted.get("document_type"), extracted_text, filename)
    extracted["document_type"] = record_type
    risk_label, risk_confidence = _predict_risk(_fill_missing_features(extracted))
    summary = extracted.get("summary") or f"Clinical document {filename} uploaded for processing."
    content = _build_readable_content(summary, record_type, risk_label, risk_confidence, extracted, filename)
    return ProcessedDocument(
        title=Path(filename).stem.replace("_", " ").strip() or "Uploaded medical document",
        content=content,
        summary=summary,
        extracted_data=extracted,
        record_type=record_type,
        risk_label=risk_label,
        risk_confidence=risk_confidence,
    )


def extract_structured_data_from_text(text: str, filename: str = "record.txt") -> dict[str, Any]:
    return _extract_structured_data_from_text(text, filename)


def fallback_extract_from_text(text: str, filename: str = "record.txt") -> dict[str, Any]:
    return _fallback_extract_from_text(text, filename)


def predict_risk_from_extracted(extracted: dict[str, Any]) -> tuple[str, str]:
    try:
        from app.services.fhir_xgboost_service import predict_fhir_xgboost_risk

        prediction = predict_fhir_xgboost_risk(extracted)
        if prediction is not None:
            return prediction.label, f"XGBoost {prediction.confidence}"
    except Exception:
        pass
    return _predict_risk(_fill_missing_features(extracted))
